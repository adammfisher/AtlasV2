#!/usr/bin/env python3
"""Compile blocks-JSON (skills/docx/schema.json) onto the .dotx style base.

All formatting flows through NAMED STYLES — Heading 1–3 (with outline levels),
Normal, Caption, Quote, List Bullet/Number, and a named table style. No direct
run formatting exists outside style definitions; the spec cannot request any
(the schema forbids inline font/size fields). Figures render real chart PNGs
via office_chart (Pillow) — never a placeholder box.
"""
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_BREAK
from docx.oxml.ns import qn
from docx.oxml.shared import OxmlElement
from docx.shared import Inches, Pt, RGBColor

import office_chart
import validate_common as vc

BRAND_DARK = "371447"   # dk1 — headings (style definitions only)
BRAND_PLUM = "650360"   # dk2 — level-2 headings
TEXT_DARK = "1B1A18"
BODY_PT = 11
FONT = "Helvetica Neue"
# named-style preference chain: first available wins (template-defined only)
TABLE_STYLES = ["Light Shading Accent 1", "Light Grid Accent 1", "Table Grid"]


def _ensure_styles(doc):
    """Define the doctrine's named styles on the document. This is the ONE
    place formatting lives; every block references a style by name."""
    styles = doc.styles
    spec = {
        "Title": dict(size=22, bold=True, color=BRAND_DARK),
        "Heading 1": dict(size=17, bold=True, color=BRAND_DARK, outline=0),
        "Heading 2": dict(size=14, bold=True, color=BRAND_PLUM, outline=1),
        "Heading 3": dict(size=12, bold=True, color=BRAND_DARK, outline=2),
        "Normal": dict(size=BODY_PT, bold=False, color=TEXT_DARK),
        "Caption": dict(size=9, bold=False, color=BRAND_PLUM, italic=True),
        "Quote": dict(size=BODY_PT, bold=False, color=BRAND_PLUM, italic=True),
    }
    for name, conf in spec.items():
        try:
            st = styles[name]
        except KeyError:
            st = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        st.font.name = FONT
        st.font.size = Pt(conf["size"])
        st.font.bold = conf["bold"]
        st.font.italic = conf.get("italic", False)
        st.font.color.rgb = RGBColor.from_string(conf["color"])
        if "outline" in conf:  # navigation pane / TOC depend on outline levels
            ppr = st.element.get_or_add_pPr()
            for old in ppr.findall(qn("w:outlineLvl")):
                ppr.remove(old)
            lvl = OxmlElement("w:outlineLvl")
            lvl.set(qn("w:val"), str(conf["outline"]))
            ppr.append(lvl)


def _table_style_name(doc) -> str:
    available = {s.name for s in doc.styles if s.type == WD_STYLE_TYPE.TABLE}
    return next((n for n in TABLE_STYLES if n in available), "Table Grid")


def _insert_toc(doc):
    """A real TOC field (levels 1–3). Word/LibreOffice populate it from the
    heading outline levels on open/update — content, not decoration."""
    doc.add_heading("Contents", level=1)
    para = doc.add_paragraph()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), r'TOC \o "1-3" \h \z \u')
    run = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = "Right-click and choose Update Field to populate the table of contents."
    run.append(text)
    fld.append(run)
    para._p.append(fld)
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def build(payload: dict, template: str, out: Path) -> dict:
    doc = Document(template)
    _ensure_styles(doc)
    for para in list(doc.paragraphs):  # template stub paragraphs
        if not para.text.strip():
            para._element.getparent().remove(para._element)

    metadata = payload.get("metadata") or {}
    doc.core_properties.title = metadata["title"]
    if metadata.get("author"):
        doc.core_properties.author = metadata["author"]
    doc.add_paragraph(metadata["title"], style="Title")

    blocks = payload["blocks"]
    heading_count = sum(1 for b in blocks if b.get("kind") == "heading")
    if heading_count >= 5:  # long-form report → real TOC from outline levels
        _insert_toc(doc)

    table_style = _table_style_name(doc)
    words = 0
    figures = 0
    tmp = Path(tempfile.mkdtemp(prefix="docxfig-"))
    for block in blocks:
        kind = block["kind"]
        if kind == "heading":
            doc.add_heading(block["text"], level=int(block["level"]))
        elif kind == "paragraph":
            doc.add_paragraph(str(block["text"]))
            words += len(str(block["text"]).split())
        elif kind in ("bulleted_list", "numbered_list"):
            style = "List Bullet" if kind == "bulleted_list" else "List Number"
            for item in block["items"]:
                doc.add_paragraph(str(item), style=style)
                words += len(str(item).split())
        elif kind == "table":
            headers = [str(h) for h in block["headers"]]
            rows = block["rows"]
            table = doc.add_table(rows=1 + len(rows), cols=len(headers))
            table.style = table_style
            for i, header in enumerate(headers):
                table.rows[0].cells[i].text = header
            for r, row in enumerate(rows):
                for c in range(len(headers)):
                    table.rows[r + 1].cells[c].text = str(row[c]) if c < len(row) else ""
        elif kind == "figure":
            figures += 1
            png = office_chart.render_chart_png(block["chart"], tmp / f"fig{figures}.png")
            doc.add_picture(str(png), width=Inches(6.0))
            doc.add_paragraph(f"Figure {figures}: {block['caption']}", style="Caption")
        elif kind == "quote":
            doc.add_paragraph(str(block["text"]), style="Quote")
            if block.get("attribution"):
                doc.add_paragraph(f"— {block['attribution']}", style="Caption")
        elif kind == "page_break":
            doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    return {
        "blocks": len(blocks),
        "words": words,
        "figures": figures,
        "bytes": out.stat().st_size,
    }


def extract_texts(path: Path) -> list:
    doc = Document(str(path))
    texts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            texts.extend(cell.text for cell in row.cells)
    return texts


def _styles_only_check(path: Path) -> dict:
    """No direct run formatting outside style definitions: document body runs
    must not carry font size/color overrides (lists/tables/headings inherit)."""
    doc = Document(str(path))
    dirty = 0
    for para in doc.paragraphs:
        for run in para.runs:
            if run.font.size is not None or run.font.color.rgb is not None:
                dirty += 1
    return vc.check("Named-styles-only (no direct run formatting)", dirty == 0)


def _hierarchy_check(path: Path) -> dict:
    doc = Document(str(path))
    level = 0
    ok = True
    for para in doc.paragraphs:
        name = para.style.name if para.style else ""
        if name.startswith("Heading "):
            lv = int(name.split()[-1])
            if level and lv > level + 1:
                ok = False
            level = lv
    return vc.check("Heading hierarchy (no level skips)", ok)


def main() -> None:
    args = vc.cli("docx builder")
    payload = vc.load_payload(args.payload)
    vc.spec_gate("docx", payload)
    out = Path(args.out)
    template = args.template or "skills/docx/templates/atlas_default.dotx"
    meta = build(payload, template, out)

    checks = [vc.openxml_audit(out), vc.zip_sanity(out)]
    reopened = Document(str(out))
    texts = extract_texts(out)
    heading_count = sum(1 for p in reopened.paragraphs if p.style.name.startswith("Heading"))
    checks.append(vc.check("Round-trip", heading_count >= 1 and any(texts)))
    checks.append(vc.placeholder_grep(texts))
    checks.append(_hierarchy_check(out))
    checks.append(_styles_only_check(out))
    checks.append(vc.soffice_convert(out, "soffice open/convert", vc.THUMBS_SKIP))
    vc.emit(out, meta, checks)


if __name__ == "__main__":
    try:
        main()
    except SystemExit:
        raise
    except Exception as err:
        vc.fail(f"{type(err).__name__}: {err}")
