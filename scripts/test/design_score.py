#!/usr/bin/env python3
"""Score a built office file with the CURRENT deterministic design gates.

Usage: design_score.py <kind> <file> [payload.json]
Emits one JSON line: {"kind", "file", "pass", "findings": [...]}.
Era-agnostic: scores files from the pre-doctrine builders (BEFORE) and the
upgraded builders (AFTER) with the same rules — that's the whole point.
"""
import json
import sys
from pathlib import Path

HERE = Path(__file__).resolve().parent
sys.path.insert(0, str(HERE.parents[1] / "scripts/office"))

import validate_common as vc  # noqa: E402


def score_pptx(path: Path, payload) -> list:
    return vc.visual_gate_pptx(path)


def score_docx(path: Path, payload) -> list:
    from docx import Document

    findings = []
    doc = Document(str(path))
    level = 0
    for para in doc.paragraphs:
        name = para.style.name if para.style else ""
        if name.startswith("Heading "):
            lv = int(name.split()[-1])
            if level and lv > level + 1:
                findings.append(f"heading level skip {level} → {lv}")
            level = lv
        for run in para.runs:
            if run.font.size is not None or run.font.color.rgb is not None:
                findings.append(f"direct run formatting on {name or 'body'} paragraph: {run.text[:30]!r}")
    texts = [p.text for p in doc.paragraphs]
    findings.extend(vc.placeholder_text_errors(texts, "docx"))
    return findings


def score_xlsx(path: Path, payload) -> list:
    from openpyxl import load_workbook

    findings = []
    recalc = vc.soffice_recalc_scan(path)
    if not recalc["ok"] and "skip" not in recalc["label"].lower():
        findings.append(recalc["label"])
    wb = load_workbook(str(path))
    for ws in wb.worksheets:
        if ws.freeze_panes != "A2":
            findings.append(f"sheet {ws.title!r}: header row not frozen")
        general_numeric = 0
        for row in ws.iter_rows():
            for cell in row:
                if isinstance(cell.value, (int, float)) and not isinstance(cell.value, bool):
                    if cell.number_format == "General":
                        general_numeric += 1
        if general_numeric:
            findings.append(f"sheet {ws.title!r}: {general_numeric} numeric cells without an explicit number format")
        if not ws.tables:
            findings.append(f"sheet {ws.title!r}: no named table style")
        if not ws.print_area:
            findings.append(f"sheet {ws.title!r}: print area not set")
    return findings


def score_pdf(path: Path, payload) -> list:
    import pdfplumber

    findings = []
    with pdfplumber.open(str(path)) as pdf:
        pages = [p.extract_text() or "" for p in pdf.pages]
    n = len(pages)
    if not all(f"of {n}" in t for t in pages):
        findings.append("running footer with page N of M missing on some pages")
    if payload:
        tables = [b for b in payload.get("sections", []) if b.get("kind") == "table"]
        for t in tables:
            leaders = [str(r[0]).strip() for r in t.get("rows", []) if r and str(r[0]).strip()]
            if leaders and not any(all(ld in page for ld in leaders) for page in pages):
                findings.append(f"table starting {leaders[0]!r} split across pages")
    return findings


def main() -> None:
    kind, file = sys.argv[1], Path(sys.argv[2])
    payload = json.loads(Path(sys.argv[3]).read_text()) if len(sys.argv) > 3 else None
    findings = {"pptx": score_pptx, "docx": score_docx, "xlsx": score_xlsx, "pdf": score_pdf}[kind](file, payload)
    print(json.dumps({"kind": kind, "file": str(file), "pass": not findings, "findings": findings[:20]}))


if __name__ == "__main__":
    main()
